import os

code = r'''# -*- coding: utf-8 -*-
import sys, os, re, shutil, tempfile
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
except ImportError:
    print("ERROR: python-docx not installed", file=sys.stderr)
    sys.exit(1)

NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def find_text(el):
    return "".join(t.text or "" for t in el.findall(f".//{{{NS}}}t"))


def clear_para(p):
    for t in p.findall(f".//{{{NS}}}t"):
        t.text = ""


def add_run(par, text, bold=False, size=10.5):
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr")
        rPr.append(OxmlElement("w:b"))
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    par.append(r)


def set_cell(cell, text, bold=False, size=10.5):
    paras = cell._element.findall(f".//{{{NS}}}p")
    if not paras:
        p = OxmlElement("w:p")
        cell._element.append(p)
        paras = [p]
    clear_para(paras[0])
    add_run(paras[0], text, bold=bold, size=size)


def parse_md(path):
    with open(path, encoding="utf-8") as f:
        text = f.read()
    d = dict(name="", contact="", intent="", summary="",
             edu_school="", edu_major="", edu_degree="", edu_dates="",
             courses="", honors="",
             projects=[], practices=[], campus=[], skills=[])
    section = None
    item = None
    for line in text.split("\n"):
        s = line.strip()
        if not s:
            continue
        if s.startswith("# "):
            d["name"] = s[2:].strip()
        elif re.match(r"^\*\*求职意向\*\*[：:]", s):
            d["intent"] = re.sub(r"\*\*求职意向\*\*[：:]", "", s).strip()
        elif re.match(r"^\*\*个人优势\*\*[：:]", s):
            d["summary"] = re.sub(r"\*\*个人优势\*\*[：:]", "", s).strip()
        elif ("\u2016" in s or "|" in s) and not s.startswith("**") and not d["contact"]:
            d["contact"] = s
        elif s == "## 教育背景":
            section = "edu"
        elif s == "## 项目经历":
            section = "projects"
        elif s == "## 实践经历":
            section = "practices"
        elif s == "## 校园经历":
            section = "campus"
        elif s == "## 综合技能":
            section = "skills"
        elif section == "edu":
            m = re.match(r"\*\*([^*]+)\*\*\s*[\u2016|]\s*([^*]+)\s*[\u2016|]\s*([^*]+)\s*[\u2016|]\s*([^*]+)\*\*", s)
            if m:
                d["edu_school"] = m.group(1).strip()
                d["edu_major"] = m.group(2).strip()
                d["edu_degree"] = m.group(3).strip()
                d["edu_dates"] = m.group(4).replace("**", "").strip()
            elif "核心课程" in s:
                d["courses"] = re.sub(r"^\*\*[^*]*\*\*[：:]\s*", "", s).strip()
            elif "荣誉奖励" in s:
                d["honors"] = re.sub(r"^\*\*[^*]*\*\*[：:]\s*", "", s).strip()
        elif s.startswith("### "):
            if section in ("projects", "practices", "campus"):
                item = dict(title=s[4:].strip(), role="", dates="", bullets=[])
                d[section].append(item)
        elif s.startswith("**") and ("\u2016" in s or "|" in s) and item is not None:
            parts = re.split(r"\*\*[^*]+\*\*\s*[\u2016|]\s*", s, maxsplit=1)
            if len(parts) >= 2:
                item["role"] = parts[0].replace("**", "").strip()
                item["dates"] = parts[1].replace("**", "").strip()
        elif s.startswith("- ") and item is not None:
            item["bullets"].append(s[2:].strip())
        elif section == "skills" and s.startswith("- "):
            d["skills"].append(s[2:].strip())
    return d


def update_header(doc, d):
    tbl = doc.tables[0]
    set_cell(tbl.rows[0].cells[0], d["name"], bold=True, size=14)
    set_cell(tbl.rows[1].cells[0], d["contact"], size=9)
    if d["intent"]:
        set_cell(tbl.rows[2].cells[0], "\u6c42\u804c\u610f\u5411\uff1a" + d["intent"], size=10)
    if d["summary"]:
        set_cell(tbl.rows[3].cells[0], "\u4e2a\u4eba\u4f18\u52bf\uff1a" + d["summary"], size=9)


def update_edu(doc, d):
    tbl = doc.tables[1]
    for row in tbl.rows[1:]:
        for c in row.cells:
            for p in c._element.findall(f".//{{{NS}}}p"):
                clear_para(p)
    cells = tbl.rows[1].cells
    if len(cells) >= 1:
        set_cell(cells[0], d["edu_school"], bold=True, size=11)
    if len(cells) >= 2:
        set_cell(cells[1], d["edu_major"], size=10.5)
    if len(cells) >= 3:
        set_cell(cells[2], d["edu_degree"], size=10.5)
    if len(cells) >= 4:
        set_cell(cells[3], d["edu_dates"], size=10.5)
    tbl2 = doc.tables[2]
    cells = tbl2.rows[0].cells
    if d["courses"] and len(cells) > 0:
        set_cell(cells[0], "\u6838\u5fc3\u8bfe\u7a0b\uff1a" + d["courses"], size=9.5)
    if d["honors"] and len(cells) > 1:
        set_cell(cells[1], "\u8363\u8a89\u5956\u52b1\uff1a" + d["honors"], size=9.5)


def update_section(doc, tidx, items, kwds):
    tbl = doc.tables[tidx]
    hi = None
    for i, row in enumerate(tbl.rows):
        for c in row.cells:
            txt = find_text(c._element)
            if any(k in txt for k in kwds):
                hi = i
                break
        if hi is not None:
            break
    if hi is None:
        return
    for row in tbl.rows[hi + 1:]:
        for c in row.cells:
            for p in c._element.findall(f".//{{{NS}}}p"):
                clear_para(p)
    ri = hi + 1
    for it in items:
        if ri >= len(tbl.rows):
            break
        row = tbl.rows[ri]
        cells = row.cells
        if len(cells) > 0:
            set_cell(cells[0], it["title"], bold=True, size=10.5)
        if len(cells) > 1 and it.get("role"):
            set_cell(cells[1], it["role"], size=9.5)
        if len(cells) > 2 and it.get("dates"):
            set_cell(cells[2], it["dates"], size=9.5)
        for bi, b in enumerate(it.get("bullets", [])):
            ci = 3 + bi
            if ci < len(cells):
                paras = cells[ci]._element.findall(f".//{{{NS}}}p")
                if paras:
                    clear_para(paras[0])
                    add_run(paras[0], "\u2022 " + b, size=9.5)
        ri += 1


def update_skills(doc, d):
    tbl = doc.tables[-1]
    for row in tbl.rows:
        for c in row.cells:
            for p in c._element.findall(f".//{{{NS}}}p"):
                clear_para(p)
    cells = tbl.rows[0].cells
    for i, s in enumerate(d.get("skills", [])):
        if i < len(cells):
            paras = cells[i]._element.findall(f".//{{{NS}}}p")
            if paras:
                clear_para(paras[0])
                add_run(paras[0], "\u2022 " + s, size=9.5)


def build(md_path, out_path, template=None):
    if not template:
        print("ERROR: --template parameter required", file=sys.stderr)
        sys.exit(1)
    if not os.path.exists(template):
        print("ERROR: template not found: " + template, file=sys.stderr)
        sys.exit(1)
    d = parse_md(md_path)
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp_path = tmp.name
    tmp.close()
    shutil.copy2(template, tmp_path)
    print("  template: " + os.path.basename(template))
    doc = Document(tmp_path)
    update_header(doc, d)
    update_edu(doc, d)
    update_section(doc, 4, d["projects"], ["\u9879\u76ee\u7ecf\u5386"])
    update_section(doc, 5, d["practices"], ["\u5b9e\u8df5\u7ecf\u5386"])
    update_section(doc, 6, d["campus"], ["\u6821\u56ed\u7ecf\u5386"])
    update_skills(doc, d)
    doc.save(out_path)
    os.unlink(tmp_path)
    print("  docx: " + os.path.basename(out_path))
    try:
        import win32com.client
        pdf_path = os.path.splitext(out_path)[0] + ".pdf"
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc_src = word.Documents.Open(os.path.abspath(out_path))
        doc_src.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc_src.Close(SaveChanges=False)
        word.Quit()
        print("  pdf: " + os.path.basename(pdf_path))
    except Exception as e:
        print("  pdf: skipped (" + str(e) + ")")
    return d


if __name__ == "__main__":
    import argparse
    p = argparse.ArgumentParser(description="Resume generator with PDF export")
    p.add_argument("md_path", help="Path to resume markdown file")
    p.add_argument("out_path", help="Output docx path")
    p.add_argument("--template", required=True, help="Template docx path")
    args = p.parse_args()
    build(args.md_path, args.out_path, template=args.template)
'''

with open(r"D:\Obsidian\SCM-Career\.tools\gen_resume.py", "w", encoding="utf-8") as f:
    f.write(code)
print("gen_resume.py rewritten: " + str(len(code.splitlines())) + " lines")
