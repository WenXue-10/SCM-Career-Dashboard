import os

code = r"""# -*- coding: utf-8 -*-
\"\"\"Resume generator with PDF export.
Usage: python -X utf8 gen_resume.py <md_path> <out_docx_path> --template <template_docx>
\"\"\"
import sys, os, re, shutil, tempfile
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
except ImportError:
    print('ERROR: python-docx not installed', file=sys.stderr)
    sys.exit(1)

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
TITLE_MAX = 15


def find_text(el):
    return ''.join(t.text or '' for t in el.findall(f'.//{{{NS}}}t'))


def clear_paragraph(p):
    for t in p.findall(f'.//{{{NS}}}t'):
        t.text = ''


def add_run(par, text, bold=False, size=10.5):
    r = OxmlElement('w:r')
    if bold:
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    par.append(r)
    return r


def set_cell_text(cell, text, bold=False, size=10.5):
    paras = cell._element.findall(f'.//{{{NS}}}p')
    if not paras:
        p = OxmlElement('w:p')
        cell._element.append(p)
        paras = [p]
    clear_paragraph(paras[0])
    add_run(paras[0], text, bold=bold, size=size)


def parse_md(md_path):
    with open(md_path, encoding='utf-8') as f:
        text = f.read()
    lines = text.split('\n')
    d = {'name': '', 'contact': '', 'intent': '', 'summary': '',
         'edu_school': '', 'edu_major': '', 'edu_degree': '', 'edu_dates': '',
         'courses': '', 'honors': '',
         'projects': [], 'practices': [], 'campus': [], 'skills': []}
    cur_section = None
    cur_item = None
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if s.startswith('# '):
            d['name'] = s[2:].strip()
        elif '\u2016' in s and not s.startswith('**') and not d['contact']:
            d['contact'] = s
        elif re.match(r'^\*\*\u6c42\u804c\u610f\u5411[：:]\*\*', s):
            d['intent'] = re.sub(r'\*\*\u6c42\u804c\u610f\u5411[：:]\*\*', '', s).strip()
        elif re.match(r'^\*\*\u4e2a\u4eba\u4f18\u52bf[：:]\*\*', s):
            d['summary'] = re.sub(r'\*\*\u4e2a\u4eba\u4f18\u52bf[：:]\*\*', '', s).strip()
        elif s == '## \u6559\u80b2\u80cc\u666f':
            cur_section = 'edu'
        elif s == '## \u9879\u76ee\u7ecf\u5386':
            cur_section = 'projects'
        elif s == '## \u5b9e\u8df5\u7ecf\u5386':
            cur_section = 'practices'
        elif s == '## \u6821\u56ed\u7ecf\u5386':
            cur_section = 'campus'
        elif s == '## \u7efc\u5408\u6280\u80fd':
            cur_section = 'skills'
        elif cur_section == 'edu':
            m = re.match(r'\*\*([^*]+)\*\*\s*[\u2016|]\s*([^*]+)\s*[\u2016|]\s*([^*]+)\s*[\u2016|]\s*([^*]+)\*\*', s)
            if m:
                d['edu_school'] = m.group(1).strip()
                d['edu_major'] = m.group(2).strip()
                d['edu_degree'] = m.group(3).strip()
                d['edu_dates'] = m.group(4).replace('**', '').strip()
            elif '\u6838\u5fc3\u8bfe\u7a0b' in s:
                d['courses'] = re.sub(r'\*\*\u6838\u5fc3\u8bfe\u7a0b[：:]\*\*', '', s).strip()
            elif '\u8363\u8a89\u5956\u52b1' in s:
                d['honors'] = re.sub(r'\*\*\u8363\u8a89\u5956\u52b1[：:]\*\*', '', s).strip()
        elif s.startswith('### '):
            if cur_section in ('projects', 'practices', 'campus'):
                cur_item = {'title': s[4:].strip(), 'role': '', 'dates': '', 'bullets': []}
                d[cur_section].append(cur_item)
        elif s.startswith('**') and ('\u2016' in s or '|' in s) and cur_item is not None:
            parts = re.split(r'\*\*[^*]+\*\*\s*[\u2016|]\s*', s, maxsplit=1)
            if len(parts) >= 2:
                cur_item['role'] = parts[0].replace('**', '').strip()
                cur_item['dates'] = parts[1].replace('**', '').strip()
            elif len(parts) == 1:
                cur_item['role'] = s.replace('**', '').strip()
        elif s.startswith('- ') and cur_item is not None:
            cur_item['bullets'].append(s[2:].strip())
        elif cur_section == 'skills' and s.startswith('- '):
            d['skills'].append(s[2:].strip())
    return d


def update_header(doc, d):
    tbl = doc.tables[0]
    set_cell_text(tbl.rows[0].cells[0], d['name'], bold=True, size=14)
    set_cell_text(tbl.rows[1].cells[0], d['contact'], size=9)
    if d['intent']:
        set_cell_text(tbl.rows[2].cells[0], '\u6c42\u804c\u610f\u5411\uff1a' + d['intent'], size=10)
    if d['summary']:
        set_cell_text(tbl.rows[3].cells[0], '\u4e2a\u4eba\u4f18\u52bf\uff1a' + d['summary'], size=9)


def update_education(doc, d):
    tbl = doc.tables[1]
    for row in tbl.rows[1:]:
        for cell in row.cells:
            for p in cell._element.findall(f'.//{{{NS}}}p'):
                clear_paragraph(p)
    row0 = tbl.rows[1]
    cells = row0.cells
    if len(cells) >= 1:
        set_cell_text(cells[0], d['edu_school'], bold=True, size=11)
    if len(cells) >= 2:
        set_cell_text(cells[1], d['edu_major'], size=10.5)
    if len(cells) >= 3:
        set_cell_text(cells[2], d['edu_degree'], size=10.5)
    if len(cells) >= 4:
        set_cell_text(cells[3], d['edu_dates'], size=10.5)
    tbl2 = doc.tables[2]
    row0 = tbl2.rows[0]
    cells = row0.cells
    if d['courses'] and len(cells) > 0:
        set_cell_text(cells[0], '\u6838\u5fc3\u8bfe\u7a0b\uff1a' + d['courses'], size=9.5)
    if d['honors'] and len(cells) > 1:
        set_cell_text(cells[1], '\u8363\u8a89\u5956\u52b1\uff1a' + d['honors'], size=9.5)


def update_section(doc, tbl_idx, items, section_keywords):
    tbl = doc.tables[tbl_idx]
    header_idx = None
    for i, row in enumerate(tbl.rows):
        for cell in row.cells:
            text = find_text(cell)
            for kw in section_keywords:
                if kw in text:
                    header_idx = i
                    break
            if header_idx is not None:
                break
        if header_idx is not None:
            break
    if header_idx is None:
        return
    for row in tbl.rows[header_idx + 1:]:
        for cell in row.cells:
            for p in cell._element.findall(f'.//{{{NS}}}p'):
                clear_paragraph(p)
    row_idx = header_idx + 1
    for item in items:
        if row_idx >= len(tbl.rows):
            break
        row = tbl.rows[row_idx]
        cells = row.cells
        if len(cells) > 0:
            set_cell_text(cells[0], item['title'], bold=True, size=10.5)
        if len(cells) > 1 and item.get('role'):
            set_cell_text(cells[1], item['role'], size=9.5)
        if len(cells) > 2 and item.get('dates'):
            set_cell_text(cells[2], item['dates'], size=9.5)
        bullet_start = 3
        for bi, bullet in enumerate(item.get('bullets', [])):
            ci = bullet_start + bi
            if ci < len(cells):
                paras = cells[ci]._element.findall(f'.//{{{NS}}}p')
                if paras:
                    clear_paragraph(paras[0])
                    add_run(paras[0], '\u2022 ' + bullet, size=9.5)
        row_idx += 1


def update_skills(doc, d):
    tbl = doc.tables[-1]
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell._element.findall(f'.//{{{NS}}}p'):
                clear_paragraph(p)
    row0 = tbl.rows[0]
    cells = row0.cells
    for i, skill in enumerate(d.get('skills', [])):
        if i < len(cells):
            paras = cells[i]._element.findall(f'.//{{{NS}}}p')
            if paras:
                clear_paragraph(paras[0])
                add_run(paras[0], '\u2022 ' + skill, size=9.5)


def build(md_path, out_path, template=None):
    if not template:
        print('ERROR: --template parameter required', file=sys.stderr)
        sys.exit(1)
    if not os.path.exists(template):
        print('ERROR: template not found: ' + template, file=sys.stderr)
        sys.exit(1)
    d = parse_md(md_path)
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp_path = tmp.name
    tmp.close()
    shutil.copy2(template, tmp_path)
    print('  template: ' + os.path.basename(template))
    doc = Document(tmp_path)
    update_header(doc, d)
    update_education(doc, d)
    update_section(doc, 4, d['projects'], ['\u9879\u76ee\u7ecf\u5386'])
    update_section(doc, 5, d['practices'], ['\u5b9e\u8df5\u7ecf\u5386'])
    update_section(doc, 6, d['campus'], ['\u6821\u56ed\u7ecf\u5386'])
    update_skills(doc, d)
    doc.save(out_path)
    os.unlink(tmp_path)
    print('  docx: ' + os.path.basename(out_path))
    try:
        import win32com.client
        pdf_path = os.path.splitext(out_path)[0] + '.pdf'
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0
        doc_src = word.Documents.Open(os.path.abspath(out_path))
        doc_src.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc_src.Close(SaveChanges=False)
        word.Quit()
        print('  pdf: ' + os.path.basename(pdf_path))
    except Exception as e:
        print('  pdf: skipped (' + str(e) + ')')
    return d


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='Resume generator with PDF export')
    parser.add_argument('md_path', help='Path to resume markdown file')
    parser.add_argument('out_path', help='Output docx path')
    parser.add_argument('--template', required=True, help='Template docx path')
    args = parser.parse_args()
    build(args.md_path, args.out_path, template=args.template)
"""

with open(r'D:\Obsidian\SCM-Career\.tools\gen_resume.py', 'w', encoding='utf-8') as f:
    f.write(code)
print('gen_resume.py written: ' + str(len(code.splitlines())) + ' lines')
