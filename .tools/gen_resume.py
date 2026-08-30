# -*- coding: utf-8 -*-
"""
简历生成脚本 v2.0
- 动态查找表格（不硬编码索引）
- 正确处理每个经历占两行的结构（标题行+要点行）
- 支持两种模板结构（8表格/7表格）
- 只调整综合技能表格行高，其他保持模板原样
"""
import sys, os, re, shutil, tempfile
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.table import WD_ROW_HEIGHT_RULE
except ImportError:
    print("ERROR: python-docx not installed", file=sys.stderr)
    sys.exit(1)

NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ========== 工具函数 ==========

def find_text(el):
    return el.text or ""

def clear_para(para):
    for run in para.runs:
        run.text = ''

def add_run(par, text, bold=False, size=10.5):
    run = par.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    # 设置英文字体为 Times New Roman，中文字体为宋体
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    return run

def set_cell(cell, text, bold=False, size=10.5):
    if cell.paragraphs:
        clear_para(cell.paragraphs[0])
        add_run(cell.paragraphs[0], text, bold=bold, size=size)
    else:
        cell.text = text

def add_formatted_text(par, text, size=9.5):
    """支持 **加粗** 格式的文本添加"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            bold_text = part[2:-2]
            add_run(par, bold_text, bold=True, size=size)
        else:
            add_run(par, part, bold=False, size=size)


# ========== 动态查找表格 ==========

def find_table_index(doc, keyword, exclude=None):
    """查找包含特定关键词的表格索引"""
    for i, tbl in enumerate(doc.tables):
        if exclude and i in exclude:
            continue
        for row in tbl.rows:
            for cell in row.cells:
                if keyword in (cell.text or ""):
                    return i
    return None

def find_all_table_indices(doc, keyword):
    """查找所有包含特定关键词的表格索引"""
    result = []
    for i, tbl in enumerate(doc.tables):
        for row in tbl.rows:
            for cell in row.cells:
                if keyword in (cell.text or ""):
                    result.append(i)
                    break
            if i in result:
                break
    return result


# ========== 解析 md ==========

def parse_md(path):
    with open(path, encoding="utf-8") as f:
        text = f.read()
    d = dict(name="", contact="", intent="", summary="",
             edu_school="", edu_major="", edu_degree="", edu_dates="",
             courses="", honors="",
             projects=[], practices=[], campus=[], skills=[])
    section = None
    item = None
    for line in text.split("\n"):
        s = line.strip()
        if not s:
            continue
        if s.startswith("# "):
            d["name"] = s[2:].strip()
        elif re.match(r"^\*\*求职意向\*\*[：:]", s):
            d["intent"] = re.sub(r"\*\*求职意向\*\*[：:]", "", s).strip()
        elif re.match(r"^\*\*个人优势\*\*[：:]", s):
            d["summary"] = re.sub(r"\*\*个人优势\*\*[：:]", "", s).strip()
        elif ("｜" in s or "|" in s) and not s.startswith("**") and not d["contact"]:
            d["contact"] = s
        elif s == "## 教育背景":
            section = "edu"
            item = None
        elif s == "## 项目经历":
            section = "projects"
            item = None
        elif s == "## 实践经历":
            section = "practices"
            item = None
        elif s == "## 校园经历":
            section = "campus"
            item = None
        elif s == "## 综合技能":
            section = "skills"
            item = None
        elif section == "edu":
            if "**" in s and ("｜" in s or "|" in s):
                parts = re.split(r"[｜|]", s)
                if len(parts) >= 2:
                    d["edu_school"] = parts[0].replace("**", "").strip()
                    d["edu_major"] = parts[1].replace("**", "").strip()
                    if len(parts) >= 3:
                        d["edu_degree"] = parts[2].replace("**", "").strip()
                    if len(parts) >= 4:
                        d["edu_dates"] = parts[3].replace("**", "").strip()
            elif "核心课程" in s:
                d["courses"] = re.sub(r"^\*\*[^*]*\*\*[：:]\s*", "", s).strip()
            elif "荣誉奖励" in s:
                d["honors"] = re.sub(r"^\*\*[^*]*\*\*[：:]\s*", "", s).strip()
        elif s.startswith("### "):
            if section in ("projects", "practices", "campus"):
                item = dict(title=s[4:].strip().replace("**", ""), role="", dates="", bullets=[])
                d[section].append(item)
        elif s.startswith("**") and ("｜" in s or "|" in s) and item is not None:
            parts = re.split(r"[｜|]", s)
            if len(parts) >= 2:
                item["role"] = parts[0].replace("**", "").strip()
                item["dates"] = parts[1].replace("**", "").strip()
        elif s.startswith("- ") and item is not None:
            item["bullets"].append(s[2:].strip())
        elif section == "skills" and s.startswith("- "):
            d["skills"].append(s[2:].strip())
    return d


# ========== 更新各板块 ==========

def update_header(doc, d):
    """更新头部：姓名、联系方式、求职意向、个人优势"""
    tbl = doc.tables[0]
    if len(tbl.rows) > 0:
        set_cell(tbl.rows[0].cells[0], d["name"], bold=True, size=22)
    if len(tbl.rows) > 1:
        set_cell(tbl.rows[1].cells[0], d["contact"], size=10.5)
    if len(tbl.rows) > 2 and d["intent"]:
        paras = list(tbl.rows[2].cells[0].paragraphs)
        if paras:
            p = paras[0]
            clear_para(p)
            add_run(p, "求职意向：", bold=True, size=10.5)
            add_run(p, d["intent"], bold=False, size=10.5)
    if len(tbl.rows) > 3 and d["summary"]:
        paras = list(tbl.rows[3].cells[0].paragraphs)
        if paras:
            p = paras[0]
            clear_para(p)
            add_run(p, "个人优势：", bold=True, size=10.5)
            add_run(p, d["summary"], bold=False, size=10.5)


def update_edu(doc, d):
    """更新教育背景：核心课程、荣誉奖励（不填学校+专业那一行，保持模板原样）"""
    # 找教育背景内容表格（包含"核心课程"的表格）
    edu_idx = find_table_index(doc, "核心课程")
    if edu_idx is None:
        return
    
    tbl = doc.tables[edu_idx]
    
    # 核心课程和荣誉奖励
    for row in tbl.rows:
        for cell in row.cells:
            text = (cell.text or "")
            if "核心课程" in text and d["courses"]:
                set_cell(cell, "核心课程：" + d["courses"], size=10.5)
            elif "荣誉奖励" in text and d["honors"]:
                set_cell(cell, "荣誉奖励：" + d["honors"], size=10.5)


def update_section(doc, items, keyword):
    """
    更新经历板块（项目/实践/校园）
    每个经历占两行：标题行（标题+角色+日期）+ 要点行
    """
    # 查找包含关键词的表格
    tbl_indices = find_all_table_indices(doc, keyword)
    if not tbl_indices:
        return
    
    # 可能有多个表格包含关键词（标题表格和内容表格），取第一个
    tbl_idx = tbl_indices[0]
    tbl = doc.tables[tbl_idx]
    
    # 找到关键词标题行的位置
    hi = None
    for i, row in enumerate(tbl.rows):
        for c in row.cells:
            txt = find_text(c._element)
            if keyword in txt:
                hi = i
                break
        if hi is not None:
            break
    if hi is None:
        return
    
    # 清空标题行之后的所有行内容
    for row in tbl.rows[hi + 1:]:
        for c in row.cells:
            for p in c.paragraphs:
                clear_para(p)
    
    # 填内容，每个经历占两行
    ri = hi + 1
    for it in items:
        if ri + 1 >= len(tbl.rows):
            break
        
        # === 第一行：标题、角色、日期 ===
        title_row = tbl.rows[ri]
        cells = title_row.cells
        
        # 标题：找第一个非空列（通常是 cells[1]）
        title_col = 1
        if len(cells) > title_col:
            paras = list(cells[title_col].paragraphs)
            if paras:
                clear_para(paras[0])
                add_formatted_text(paras[0], it["title"], size=10.5)
        
        # 角色和日期：根据列数动态找
        # 常见结构：cells[1]=标题, cells[2/3]=角色, cells[3/4]=日期
        role_col = None
        date_col = None
        if len(cells) >= 6:
            # 6列结构：cells[1]=标题, cells[3]=角色, cells[4]=日期
            role_col = 3
            date_col = 4
        elif len(cells) >= 5:
            # 5列结构：cells[1]=标题, cells[3]=角色, cells[4]=日期
            role_col = 3
            date_col = 4
        elif len(cells) >= 4:
            # 4列结构：cells[1]=标题, cells[2]=角色, cells[3]=日期
            role_col = 2
            date_col = 3
        
        if role_col is not None and len(cells) > role_col and it.get("role"):
            paras = list(cells[role_col].paragraphs)
            if paras:
                clear_para(paras[0])
                add_formatted_text(paras[0], it["role"], size=10.5)
        
        if date_col is not None and len(cells) > date_col and it.get("dates"):
            paras = list(cells[date_col].paragraphs)
            if paras:
                clear_para(paras[0])
                add_formatted_text(paras[0], it["dates"], size=10.5)
        
        # === 第二行：要点 ===
        bullet_row = tbl.rows[ri + 1]
        cells = bullet_row.cells
        
        # 要点通常在 cells[1]（合并单元格）
        bullet_col = 1
        if len(cells) > bullet_col:
            cell = cells[bullet_col]
            # 清空所有段落
            for p in cell.paragraphs:
                clear_para(p)
            # 填要点，每个要点一段
            for bi, b in enumerate(it.get("bullets", [])):
                if bi == 0:
                    paras = list(cell.paragraphs)
                    if paras:
                        add_formatted_text(paras[0], "\u2022 " + b, size=10.5)
                else:
                    p = cell.add_paragraph()
                    add_formatted_text(p, "\u2022 " + b, size=10.5)
        
        ri += 2  # 每个经历占两行


def update_skills(doc, d):
    """更新综合技能：所有技能填到内容表格的格[1]，每个一段（和模板一致）"""
    # 综合技能内容表格：最后一个表格
    tbl = doc.tables[-1]
    
    # 清空所有单元格
    for row in tbl.rows:
        for c in row.cells:
            for p in c.paragraphs:
                clear_para(p)
    
    skills = d.get("skills", [])
    if not skills:
        _adjust_skills_height(tbl, 0)
        return
    
    # 模板结构：格[0]空，格[1]放所有技能（每个一段）
    cells = tbl.rows[0].cells
    # 找有内容的单元格（通常是格[1]），如果都空就用最后一个
    target_cell = None
    for ci in range(len(cells) - 1, -1, -1):
        target_cell = cells[ci]
        break
    
    if target_cell is None:
        return
    
    # 先删除所有多余段落，只保留一个
    while len(target_cell.paragraphs) > 1:
        p = target_cell.paragraphs[-1]
        p._element.getparent().remove(p._element)
    
    # 填所有技能，每个一段
    for i, s in enumerate(skills):
        if i == 0:
            paras = list(target_cell.paragraphs)
            if paras:
                clear_para(paras[0])
                add_formatted_text(paras[0], s, size=10.5)
        else:
            p = target_cell.add_paragraph()
            add_formatted_text(p, s, size=10.5)
    
    # 动态调整综合技能表格行高
    _adjust_skills_height(tbl, len(skills))


def _adjust_skills_height(tbl, skills_count):
    """只调整综合技能表格的行高，根据技能条数动态调整"""
    if skills_count <= 2:
        h = Pt(28)
    elif skills_count <= 4:
        h = Pt(42)
    else:
        h = Pt(56)
    for row in tbl.rows:
        row.height = h
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


# ========== 主函数 ==========

def build(md_path, out_path, template=None):
    if not template:
        print("ERROR: --template parameter required", file=sys.stderr)
        sys.exit(1)
    if not os.path.exists(template):
        print("ERROR: template not found: " + template, file=sys.stderr)
        sys.exit(1)
    
    d = parse_md(md_path)
    
    # 复制模板
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp_path = tmp.name
    tmp.close()
    shutil.copy2(template, tmp_path)
    print("  template: " + os.path.basename(template))
    
    doc = Document(tmp_path)
    
    # 更新各板块（动态查找表格，不硬编码索引）
    update_header(doc, d)
    update_edu(doc, d)
    update_section(doc, d["projects"], "项目经历")
    update_section(doc, d["practices"], "实践经历")
    update_section(doc, d["campus"], "校园经历")
    update_skills(doc, d)
    
    doc.save(out_path)
    os.unlink(tmp_path)
    print("  docx: " + os.path.basename(out_path))
    return d


if __name__ == "__main__":
    import argparse
    p = argparse.ArgumentParser(description="Resume generator v2.0")
    p.add_argument("md_path", help="Path to resume markdown file")
    p.add_argument("out_path", help="Output docx path")
    p.add_argument("--template", required=True, help="Template docx path")
    args = p.parse_args()
    build(args.md_path, args.out_path, template=args.template)
