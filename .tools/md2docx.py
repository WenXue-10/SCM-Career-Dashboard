#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
面试备考资料 Markdown 转 Word v2.0（ins风）
- 封面大字体
- 自动生成目录（二级+三级标题，带超链接可跳转）
- 字体统一（中文宋体/英文Times New Roman，标题层级清晰）
- ins风主题（柔和配色、清新排版、适当留白）
"""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ========== ins风配色 ==========
COLOR_PRIMARY = RGBColor(0x4A, 0x6F, 0x5A)      # 莫兰迪绿（主色）
COLOR_SECONDARY = RGBColor(0x8B, 0xA8, 0x9A)    # 浅绿（辅助）
COLOR_ACCENT = RGBColor(0xC4, 0x9A, 0x6A)       # 奶茶色（强调）
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)          # 深灰（正文）
COLOR_LIGHT = RGBColor(0x99, 0x99, 0x99)         # 浅灰（辅助文字）
COLOR_BG = RGBColor(0xF5, 0xF2, 0xEC)            # 米白（背景感）

CN_FONT = "宋体"
EN_FONT = "Times New Roman"
CN_FONT_TITLE = "微软雅黑"  # 标题用微软雅黑，更ins风


def set_run_font(run, size=12, bold=False, color=None, cn_font=CN_FONT):
    """统一设置字体"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = EN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), EN_FONT)
    rfonts.set(qn('w:hAnsi'), EN_FONT)
    rfonts.set(qn('w:eastAsia'), cn_font)
    if color:
        run.font.color.rgb = color


def add_bookmark(paragraph, bookmark_name):
    """添加书签"""
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), '1')
    bookmark_start.set(qn('w:name'), bookmark_name)
    paragraph._p.append(bookmark_start)
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), '1')
    paragraph._p.append(bookmark_end)


def add_hyperlink(paragraph, text, bookmark_name, size=11, color=COLOR_PRIMARY):
    """添加指向书签的超链接"""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # 字体
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), EN_FONT)
    rFonts.set(qn('w:hAnsi'), EN_FONT)
    rFonts.set(qn('w:eastAsia'), CN_FONT)
    rPr.append(rFonts)
    
    # 字号
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(size * 2))
    rPr.append(szCs)
    
    # 颜色
    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), '%02X%02X%02X' % (color[0], color[1], color[2]))
    rPr.append(color_elem)
    
    # 下划线
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    run.append(rPr)
    
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number_footer(section):
    """添加页码页脚"""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— ")
    set_run_font(run, 9, color=COLOR_LIGHT)
    fld1 = OxmlElement('w:fldSimple')
    fld1.set(qn('w:instr'), 'PAGE')
    run._element.addnext(fld1)
    r2 = p.add_run(" —")
    set_run_font(r2, 9, color=COLOR_LIGHT)


def add_header(section, text):
    """添加页眉"""
    hdr = section.header
    p = hdr.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 9, color=COLOR_LIGHT, cn_font=CN_FONT_TITLE)


def parse_inline(paragraph, text, base_size=11, base_color=COLOR_TEXT):
    """解析行内格式（**加粗**），统一字体"""
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, seg in enumerate(parts):
        if not seg:
            continue
        run = paragraph.add_run(seg)
        set_run_font(run, base_size, bold=(i % 2 == 1), color=base_color)


def add_heading_with_bookmark(doc, text, level, bookmark_id):
    """添加带书签的标题"""
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(18 if level == 1 else (12 if level == 2 else 8))
    h.paragraph_format.space_after = Pt(8 if level == 1 else (6 if level == 2 else 4))
    
    bookmark_name = f"heading_{bookmark_id}"
    add_bookmark(h, bookmark_name)
    
    if level == 1:
        # 二级标题（##）：ins风，主色
        run = h.add_run(text)
        set_run_font(run, 16, bold=True, color=COLOR_PRIMARY, cn_font=CN_FONT_TITLE)
    elif level == 2:
        # 三级标题（###）：辅助色
        run = h.add_run(text)
        set_run_font(run, 13, bold=True, color=COLOR_SECONDARY, cn_font=CN_FONT_TITLE)
    elif level == 3:
        # 四级标题（####）：具体问题，强调色
        run = h.add_run(text)
        set_run_font(run, 11.5, bold=True, color=COLOR_ACCENT, cn_font=CN_FONT_TITLE)
    else:
        run = h.add_run(text)
        set_run_font(run, 11, bold=True, color=COLOR_TEXT, cn_font=CN_FONT_TITLE)
    
    return bookmark_name


def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    
    # 基础样式
    style = doc.styles['Normal']
    style.font.name = EN_FONT
    style.font.size = Pt(11)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), CN_FONT)
    pf = style.paragraph_format
    pf.line_spacing = 1.6
    pf.space_after = Pt(4)
    
    # 页面边距（ins风，更宽松）
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # 提取标题和日期
    title = "面试备考资料"
    date = ""
    company = ""
    position = ""
    for ln in lines:
        if ln.startswith('# '):
            title = ln[2:].strip()
            break
    for ln in lines:
        m = re.search(r'生成日期\s*[:：]\s*([\d\-]+)', ln)
        if m:
            date = m.group(1)
            break
    for ln in lines:
        m = re.search(r'公司\s*[:：]\s*(.+)', ln)
        if m:
            company = m.group(1).strip()
            break
    for ln in lines:
        m = re.search(r'岗位\s*[:：]\s*(.+)', ln)
        if m:
            position = m.group(1).strip()
            break

    # ========== 封面页（ins风大字体） ==========
    sec = doc.sections[0]
    add_header(sec, "面试备考资料 · 内部使用")
    add_page_number_footer(sec)
    
    # 顶部留白
    for _ in range(4):
        doc.add_paragraph()
    
    # 主标题（大字体）
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(20)
    cr = sp.add_run(title)
    set_run_font(cr, 28, bold=True, color=COLOR_PRIMARY, cn_font=CN_FONT_TITLE)
    
    # 分隔线（用文字模拟）
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_p.paragraph_format.space_after = Pt(20)
    line_run = line_p.add_run("— ✦ —")
    set_run_font(line_run, 14, color=COLOR_ACCENT)
    
    # 副标题
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(12)
    sr = sub.add_run("面试备考资料")
    set_run_font(sr, 16, color=COLOR_SECONDARY, cn_font=CN_FONT_TITLE)
    
    # 公司和岗位
    if company:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(6)
        cpr = cp.add_run(company)
        set_run_font(cpr, 13, color=COLOR_TEXT, cn_font=CN_FONT_TITLE)
    
    if position:
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_after = Pt(30)
        ppr = pp.add_run(position)
        set_run_font(ppr, 12, color=COLOR_LIGHT)
    
    # 个人信息
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_after = Pt(8)
    info_run = info_p.add_run("文雪 · 山东大学 · 供应链管理 · 2027届")
    set_run_font(info_run, 11, color=COLOR_LIGHT)
    
    if date:
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dr = dp.add_run(f"生成日期：{date}")
        set_run_font(dr, 10, color=COLOR_LIGHT)
    
    doc.add_page_break()

    # ========== 第一遍：收集所有标题（用于生成目录） ==========
    headings = []  # [(level, text, bookmark_id)]
    bookmark_counter = 0
    i = 0
    n = len(lines)
    in_toc = False
    while i < n:
        line = lines[i]
        # 跳过目录部分（包括其下面的子标题）
        if line.startswith('## 目录'):
            in_toc = True
            i += 1
            continue
        if in_toc:
            if line.startswith('## '):
                in_toc = False
            else:
                i += 1
                continue
        
        if line.startswith('## '):
            text = line[3:].strip()
            bookmark_counter += 1
            headings.append((1, text, bookmark_counter))
        elif line.startswith('### '):
            text = line[4:].strip()
            bookmark_counter += 1
            headings.append((2, text, bookmark_counter))
        elif line.startswith('#### '):
            text = line[5:].strip()
            bookmark_counter += 1
            headings.append((3, text, bookmark_counter))
        i += 1

    # ========== 目录页 ==========
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.space_after = Pt(20)
    ttr = toc_title.add_run("目  录")
    set_run_font(ttr, 20, bold=True, color=COLOR_PRIMARY, cn_font=CN_FONT_TITLE)
    
    # 目录分隔线
    toc_line = doc.add_paragraph()
    toc_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_line.paragraph_format.space_after = Pt(16)
    tlr = toc_line.add_run("— ✦ —")
    set_run_font(tlr, 12, color=COLOR_ACCENT)
    
    # 目录内容（带超链接）
    current_h1 = None
    for level, text, bid in headings:
        bookmark_name = f"heading_{bid}"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.4
        
        if level == 1:
            # 二级标题（##）：加粗，无缩进
            p.paragraph_format.left_indent = Cm(0)
            add_hyperlink(p, text, bookmark_name, size=11.5, color=COLOR_PRIMARY)
            current_h1 = text
        elif level == 2:
            # 三级标题（###）：一级缩进
            p.paragraph_format.left_indent = Cm(0.8)
            add_hyperlink(p, text, bookmark_name, size=10.5, color=COLOR_SECONDARY)
        else:
            # 四级标题（####）：二级缩进，更小字号
            p.paragraph_format.left_indent = Cm(1.6)
            add_hyperlink(p, text, bookmark_name, size=9.5, color=COLOR_LIGHT)
    
    doc.add_page_break()

    # ========== 第二遍：生成正文内容 ==========
    i = 0
    bookmark_counter = 0
    in_table = False
    frontmatter_skipped = False
    while i < n:
        line = lines[i]
        
        # 跳过 frontmatter（只在文件开头跳过一次）
        if not frontmatter_skipped and line.strip() == '---':
            frontmatter_skipped = True
            i += 1
            while i < n and lines[i].strip() != '---':
                i += 1
            i += 1
            continue
        
        # 跳过标题行（已在封面）
        if line.startswith('# '):
            i += 1
            continue
        
        # 跳过手动写的目录部分
        if line.startswith('## 目录'):
            i += 1
            while i < n and not lines[i].startswith('## '):
                i += 1
            continue
        
        # 二级标题（##）
        if line.startswith('## '):
            text = line[3:].strip()
            bookmark_counter += 1
            add_heading_with_bookmark(doc, text, 1, bookmark_counter)
            i += 1
            continue
        
        # 三级标题（###）
        if line.startswith('### '):
            text = line[4:].strip()
            bookmark_counter += 1
            add_heading_with_bookmark(doc, text, 2, bookmark_counter)
            i += 1
            continue
        
        # 四级标题（####）
        if line.startswith('#### '):
            text = line[5:].strip()
            bookmark_counter += 1
            add_heading_with_bookmark(doc, text, 3, bookmark_counter)
            i += 1
            continue
        
        # 表格
        if line.strip().startswith('|'):
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(cells)
                i += 1
            # 清理分隔行
            clean = []
            for r in rows:
                if all(re.fullmatch(r':?-+:?', c) for c in r if c != ''):
                    continue
                clean.append(r)
            if clean:
                t = doc.add_table(rows=len(clean), cols=max(len(r) for r in clean))
                t.style = 'Light Grid Accent 1'
                for ri, r in enumerate(clean):
                    for ci, c in enumerate(r):
                        cell = t.cell(ri, ci)
                        cell.text = ''
                        parse_inline(cell.paragraphs[0], c, 10)
                doc.add_paragraph()  # 表格后空行
            continue
        
        # 引用块
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # 引用块用浅色背景感（通过左侧边框模拟）
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '8')
            left.set(qn('w:color'), '%02X%02X%02X' % (COLOR_ACCENT[0], COLOR_ACCENT[1], COLOR_ACCENT[2]))
            pBdr.append(left)
            pPr.append(pBdr)
            parse_inline(p, line[2:].strip(), 10.5, COLOR_LIGHT)
            i += 1
            continue
        
        # 分隔线
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            r = p.add_run("— · —")
            set_run_font(r, 10, color=COLOR_LIGHT)
            i += 1
            continue
        
        # 空行
        if line.strip() == '':
            i += 1
            continue
        
        # 普通正文
        p = doc.add_paragraph()
        parse_inline(p, line.strip(), 11, COLOR_TEXT)
        i += 1

    doc.save(docx_path)
    print(f"✅ 已生成：{docx_path}")
    print(f"   目录包含 {len(headings)} 个标题（带超链接可跳转）")


if __name__ == '__main__':
    for pair in sys.argv[1:]:
        md, docx = pair.split('::')
        convert(md, docx)
