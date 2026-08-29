# -*- coding: utf-8 -*-
"""编码与 docx 辅助函数，供 .workbuddy/ 下所有脚本共用。"""
import sys
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

def read_docx(path):
    """读取 docx 并返回 Document 对象。"""
    return Document(path)

def docx_to_text(doc):
    """提取 docx 所有段落文本列表。"""
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def docx_tables_text(doc):
    """提取 docx 所有表格单元格文本，按表格顺序返回二维列表。"""
    result = []
    for tbl in doc.tables:
        tbl_data = []
        for row in tbl.rows:
            row_data = [c.text.strip() for c in row.cells]
            tbl_data.append(row_data)
        result.append(tbl_data)
    return result
