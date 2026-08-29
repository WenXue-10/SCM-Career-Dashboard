#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将简历 Markdown 转为一页式规范 Word（无封面/目录，页眉简历名，页脚页码，宋体/Times，紧凑排版）。
用法: python resume2docx.py "输入.md::输出.docx" ["输入2.md::输出2.docx" ...]
"""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CN_FONT = "宋体"
EN_FONT = "Times New Roman"
GRAY = RGBColor(0x60, 0x60, 0x60)

def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = EN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), EN_FONT)
    rfonts.set(qn('w:hAnsi'), EN_FONT)
    rfonts.set(qn('w:eastAsia'), CN_FONT)
    if color:
        run.font.color.rgb = color

def add_inline(par, text, size=10.5, bold_color=None):
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, seg in enumerate(parts):
        if not seg:
            continue
        run = par.add_run(seg)
        set_run_font(run, size, bold=(i % 2 == 1), color=bold_color if i % 2 == 1 else None)

def add_bottom_border(par):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '404040')
    pbdr.append(bottom)
    pPr.append(pbdr)

def add_page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("第 ")
    set_run_font(r1, 8, color=GRAY)
    fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE')
    r1._element.addnext(fld1)
    r2 = p.add_run(" 页 / 共 ")
    set_run_font(r2, 8, color=GRAY)
    fld2 = OxmlElement('w:fldSimple'); fld2.set(qn('w:instr'), 'NUMPAGES')
    r2._element.addnext(fld2)
    r3 = p.add_run(" 页")
    set_run_font(r3, 8, color=GRAY)

def add_header(section, text):
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 8, color=GRAY)

def tight(par, before=0, after=0, line=1.15):
    pf = par.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line

def md_to_docx(md_path, docx_path, resume_name):
    with open(md_path, encoding='utf-8') as f:
        lines = [l.rstrip() for l in f.read().split('\n')]

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.2); sec.bottom_margin = Cm(1.2)
    sec.left_margin = Cm(1.6); sec.right_margin = Cm(1.6)
    add_header(sec, resume_name)
    add_page_number_footer(sec)

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = EN_FONT
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), CN_FONT)

    i, n = 0, len(lines)
    name_done = False
    while i < n:
        line = lines[i].strip()
        if not line:
            i += 1; continue
        # 姓名标题
        m = re.match(r'^#\s+(.+)$', line)
        if m and not name_done:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tight(p, 0, 2, 1.0)
            run = p.add_run(m.group(1))
            set_run_font(run, 17, bold=True)
            name_done = True
            i += 1; continue
        # 联系方式行（姓名后第一个非空行且含｜或电话特征）
        if name_done and ('｜' in line) and re.search(r'\d{6,}|@', line) and not line.startswith('**'):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tight(p, 0, 3, 1.0)
            add_inline(p, line, size=9.5)
            i += 1; continue
        # 二级标题（板块）
        m = re.match(r'^##\s+(.+)$', line)
        if m:
            p = doc.add_paragraph()
            tight(p, 5, 2, 1.0)
            run = p.add_run(m.group(1))
            set_run_font(run, 12, bold=True)
            add_bottom_border(p)
            i += 1; continue
        # 三级标题（经历名）
        m = re.match(r'^###\s+(.+)$', line)
        if m:
            p = doc.add_paragraph()
            tight(p, 3, 0, 1.0)
            add_inline(p, m.group(1), size=11)
            i += 1; continue
        # 角色行（**角色** ｜ 时间）
        if line.startswith('**') and '｜' in line:
            p = doc.add_paragraph()
            tight(p, 0, 1, 1.0)
            add_inline(p, line, size=9.5, bold_color=GRAY)
            i += 1; continue
        # 列表项
        m = re.match(r'^-\s+(.+)$', line)
        if m:
            p = doc.add_paragraph()
            tight(p, 0, 1, 1.15)
            p.paragraph_format.left_indent = Cm(0.5)
            bullet = p.add_run("• ")
            set_run_font(bullet, 10.5)
            add_inline(p, m.group(1), size=10.5)
            i += 1; continue
        # 求职意向/个人优势标签段
        p = doc.add_paragraph()
        tight(p, 0, 1, 1.15)
        add_inline(p, line, size=10.5)
        i += 1

    doc.save(docx_path)
    print('OK:', docx_path)

if __name__ == '__main__':
    for arg in sys.argv[1:]:
        md, docx = arg.split('::')
        nm = re.sub(r'\.(md|docx)$', '', md.split('/')[-1].split('\\')[-1])
        md_to_docx(md, docx, nm)
