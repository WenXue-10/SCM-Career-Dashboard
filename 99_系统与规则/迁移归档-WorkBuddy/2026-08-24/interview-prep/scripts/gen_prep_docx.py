# -*- coding: utf-8 -*-
"""
面试/笔试备考资料 -> 精美 Word（ins 风 + 两级可点击目录）。
用法: python gen_prep_docx.py <输入.md> <输出.docx>
md 头部为元数据行(KEY: VALUE)，之后为正文。
"""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MAIN = RGBColor(0xB5, 0x83, 0x8D)
SUB = RGBColor(0x6D, 0x9A, 0x8E)
BODY = RGBColor(0x3F, 0x3A, 0x3A)
GREY = RGBColor(0x80, 0x80, 0x80)
TBL_HDR_FILL = "B5838D"
TBL_HDR_TXT = RGBColor(0xFF, 0xFF, 0xFF)
TBL_ZEBRA = "FBF4F1"
TBL_ZEBRA_ALT = "FFFFFF"
TBL_BORDER = "D8C3BC"
HR = "E0C3B5"
TOC_LINE = "E0C3B5"

HL = {
    '📌': "EAF2EC", '🔍': "EAF0F5", '📝': "FBF3D9", '💡': "FBF3D9",
    '📖': "EAF2EC", '📚': "EAF2EC", '🎯': "E3F0E8", '📎': "F5EFE6",
    '⚠️': "FBF3D9", '✅': "E3F0E8", '❌': "FBE3E6", '📋': "EAF0F5",
    '💼': "EAF2EC", '📑': "EAF0F5",
}
CN_FONT = '微软雅黑'


def set_run_font(run, cn=CN_FONT, en='Calibri', size=10.5, bold=False, color=BODY):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), cn)
    rfonts.set(qn('w:ascii'), en)
    rfonts.set(qn('w:hAnsi'), en)


def shade_para(p, color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color)
    pPr.append(shd)


def shade_cell(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def add_inline(paragraph, text, size=10.5, color=BODY):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for p in parts:
        if not p:
            continue
        if p.startswith('**') and p.endswith('**'):
            run = paragraph.add_run(p[2:-2])
            run.font.bold = True
        else:
            run = paragraph.add_run(p)
        set_run_font(run, size=size, color=color)


def add_bookmark(paragraph, name, bmid):
    run = paragraph.add_run()
    bs = OxmlElement('w:bookmarkStart')
    bs.set(qn('w:id'), str(bmid)); bs.set(qn('w:name'), name)
    run._r.append(bs)
    run2 = paragraph.add_run()
    be = OxmlElement('w:bookmarkEnd')
    be.set(qn('w:id'), str(bmid))
    run2._r.append(be)


def add_hyperlink(paragraph, anchor, text, color=MAIN, size=11, bold=False):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle'); rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '%02X%02X%02X' % (color[0], color[1], color[2]))
    rPr.append(color_el)
    if bold:
        b = OxmlElement('w:b'); rPr.append(b)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_h1(doc, text, bmid, anchor):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    add_bookmark(p, anchor, bmid)
    run = p.add_run(text)
    set_run_font(run, size=18, bold=True, color=MAIN)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), 'B5838D')
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def add_h2(doc, text, bmid, anchor):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    add_bookmark(p, anchor, bmid)
    run = p.add_run(text)
    set_run_font(run, size=15, bold=True, color=SUB)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=BODY)
    return p


def render_normal(doc, text, indent=False, bullet=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    if bullet:
        text = re.sub(r'^-\s+', '• ', text)
    add_inline(p, text)
    return p


def render_highlight(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    add_inline(p, text)
    for em, col in HL.items():
        if text.startswith(em):
            shade_para(p, col)
            break
    return p


def render_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.9)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    add_inline(p, text)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '6'); left.set(qn('w:color'), HR)
    pbdr.append(left); pPr.append(pbdr)
    return p


def render_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), HR)
    pbdr.append(bottom); pPr.append(pbdr)


def set_table_borders(table, color="D8C3BC", sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)


def set_col_widths(table, header, doc):
    sec = doc.sections[0]
    usable = sec.page_width - sec.left_margin - sec.right_margin
    ncol = len(header)
    weights = [1.3] * ncol
    wide_terms = ['回答', '要点', '预案', '建议', '说明', '追问', '内容', '要点与示例', '解析', '回答示例']
    for j, h in enumerate(header):
        for t in wide_terms:
            if t in h:
                weights[j] = 3.2
                break
    total = sum(weights)
    for j, col in enumerate(table.columns):
        w = int(usable * (weights[j] / total))
        col.width = w
        for cell in col.cells:
            cell.width = w


def render_table(doc, rows):
    data = []
    for r in rows:
        cells = [c.strip() for c in r.strip().strip('|').split('|')]
        data.append(cells)
    header = data[0]
    if len(data) > 1 and all(c and set(c) <= set('-: ') for c in data[1]):
        body = data[2:]
    else:
        body = data[1:]
    ncol = len(header)
    table = doc.add_table(rows=1, cols=ncol)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    set_table_borders(table, TBL_BORDER, 4)
    hdr = table.rows[0].cells
    for j, c in enumerate(header):
        hdr[j].text = ''
        para = hdr[j].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(para, c, size=10.5, color=TBL_HDR_TXT)
        for run in para.runs:
            run.font.bold = True
        shade_cell(hdr[j], TBL_HDR_FILL)
    for ri, row in enumerate(body):
        cells = table.add_row().cells
        zebra = TBL_ZEBRA if ri % 2 == 0 else TBL_ZEBRA_ALT
        for j in range(ncol):
            val = row[j] if j < len(row) else ''
            cells[j].text = ''
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, val, size=10)
            shade_cell(cells[j], zebra)
    set_col_widths(table, header, doc)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)


def build_cover(doc, meta):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(meta.get('company', '')), size=22, bold=True, color=MAIN)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(meta.get('title', '')), size=15, bold=True, color=SUB)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(meta.get('role', '')), size=11, color=BODY)
    render_hr(doc)
    for label in ['name', 'date', 'source']:
        if meta.get(label):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(meta[label]), size=10.5, color=BODY)
    doc.add_page_break()


def build_toc(doc, toc_items):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run('📑 目 录'), size=15, bold=True, color=MAIN)
    p.paragraph_format.space_after = Pt(10)
    for (text, anchor, level) in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6 + 0.9 * (level - 1))
        p.paragraph_format.space_after = Pt(4 if level == 1 else 2)
        add_hyperlink(p, anchor, text, color=MAIN if level == 1 else SUB, size=11.5 if level == 1 else 10.5, bold=(level == 1))
        if level == 1:
            dot = p.add_run(' ' + '·' * 5)
            set_run_font(dot, size=10, color=RGBColor.from_string(TOC_LINE))
    doc.add_page_break()


def main():
    if len(sys.argv) < 3:
        print("用法: python gen_prep_docx.py <输入.md> <输出.docx>")
        sys.exit(1)
    SRC, OUT = sys.argv[1], sys.argv[2]
    with open(SRC, encoding='utf-8') as f:
        lines = f.read().split('\n')

    # 头部元数据
    meta = {}
    body_start = 0
    for i, ln in enumerate(lines):
        s = ln.strip()
        if not s:
            body_start = i + 1
            continue
        if re.match(r'^[A-Za-z_]+:\s*', s) and not s.startswith('#'):
            k, v = s.split(':', 1)
            meta[k.strip()] = v.strip()
            body_start = i + 1
        else:
            break

    # 收集目录（一级 # 与二级 ##）
    toc = []
    for ln in lines[body_start:]:
        s = ln.strip()
        m1 = re.match(r'^#\s+(.+)$', s)
        if m1:
            toc.append((m1.group(1).strip(), 'sec_%d' % len(toc), 1))
            continue
        m2 = re.match(r'^##\s+(.+)$', s)
        if m2:
            toc.append((m2.group(1).strip(), 'sec_%d' % len(toc), 2))

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(3.0)
    sec.different_first_page_header_footer = True
    hdr = sec.first_page_header.paragraphs[0]
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(hdr.add_run('📋 文雪求职备考资料・内部使用'), size=9, color=GREY)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(10.5)
    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), CN_FONT)
    style.paragraph_format.line_spacing = 1.25

    build_cover(doc, meta)
    build_toc(doc, toc)

    anchor_map = {t: a for t, a, l in toc}
    bmid = 0
    i = body_start
    n = len(lines)
    while i < n:
        s = lines[i].strip()
        if s.startswith('|'):
            tbl = []
            while i < n and lines[i].strip().startswith('|'):
                tbl.append(lines[i].strip()); i += 1
            render_table(doc, tbl); continue
        if not s:
            i += 1; continue
        if s.startswith('### '):
            add_h3(doc, s[4:]); i += 1; continue
        if s.startswith('# '):
            t = s[2:].strip()
            add_h1(doc, t, bmid, anchor_map.get(t, 'sec_x')); bmid += 1
            i += 1; continue
        if s.startswith('## '):
            t = s[3:]
            add_h2(doc, t, bmid, anchor_map.get(t, 'sec_x')); bmid += 1
            i += 1; continue
        if s == '---':
            render_hr(doc); i += 1; continue
        if s.startswith('>'):
            render_quote(doc, s[1:].lstrip()); i += 1; continue
        matched = False
        for em in HL:
            if s.startswith(em):
                render_highlight(doc, s); matched = True; break
        if matched:
            i += 1; continue
        if re.match(r'^\d+\.\s', s):
            render_normal(doc, s, indent=True); i += 1; continue
        if s.startswith('- '):
            render_normal(doc, s, indent=True, bullet=True); i += 1; continue
        render_normal(doc, s); i += 1

    doc.save(OUT)
    print("SAVED:", OUT, "| 目录条目:", len(toc))


if __name__ == '__main__':
    main()
