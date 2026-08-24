# -*- coding: utf-8 -*-
"""
求职背调报告 -> 精美 Word（ins 风 / 可爱主题 / 可点击目录）。

用法:
    python gen_report_docx.py <输入报告预览.md> <输出.docx>

特性:
    - 封面页（微软雅黑 + 干玫瑰粉主色）
    - 目录页（收录所有一级标题，内部超链接点击跳转）
    - 正文：一级/二级标题（加粗加大）、表格（精致风：实心玫瑰表头白字 + 斑马纹 + 柔和边框 + 加宽列宽）、📌⚠️✅❌ 高亮块、分隔线
    - 段落间距收紧（行距 1.25，标题段前/段后更小）
    - 柔和低饱和莫兰迪配色
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- ins 风配色 ----------
MAIN = RGBColor(0xB5, 0x83, 0x8D)      # 干玫瑰粉（一级标题/封面）
SUB = RGBColor(0x6D, 0x9A, 0x8E)       # 薄荷绿灰（二级标题）
BODY = RGBColor(0x3F, 0x3A, 0x3A)      # 暖灰黑（正文）
GREY = RGBColor(0x80, 0x80, 0x80)      # 页眉灰
HDR_FILL = "F3E3DC"                     # 表头浅奶粉（备用）
HDR_TXT = RGBColor(0x7A, 0x5C, 0x61)    # 表头深玫瑰字（备用）
# 精致风表格配色
TBL_HDR_FILL = "B5838D"                 # 实心干玫瑰粉表头
TBL_HDR_TXT = RGBColor(0xFF, 0xFF, 0xFF)  # 表头白字
TBL_ZEBRA = "FBF4F1"                    # 斑马纹浅玫瑰
TBL_ZEBRA_ALT = "FFFFFF"                # 斑马纹白
TBL_BORDER = "D8C3BC"                   # 柔和边框玫瑰灰
HR = "E0C3B5"                           # 分隔线浅陶土
TOC_LINE = "E0C3B5"

HL = {
    '📌': "EAF2EC",   # 浅灰绿
    '⚠️': "FBF3D9",   # 浅鹅黄
    '✅': "E3F0E8",   # 浅薄荷
    '❌': "FBE3E6",   # 浅粉红
}

# 一级标题对应图标（用于目录）
SECTION_ICON = {
    '第一部分': '💼', '第二部分': '📋', '第三部分': '🔍',
    '第四部分': '💡', '第五部分': '📚', '第六部分': '🎯',
    '信息缺口': '⚠️',
}

CN_FONT = '微软雅黑'


def set_run_font(run, cn=CN_FONT, en='Calibri', size=10.5, bold=False, color=BODY):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), cn)
    rfonts.set(qn('w:ascii'), en)
    rfonts.set(qn('w:hAnsi'), en)


def shade_para(p, color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color)
    pPr.append(shd)


def shade_cell(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def add_inline(paragraph, text, size=10.5, color=BODY):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for p in parts:
        if not p:
            continue
        if p.startswith('**') and p.endswith('**'):
            run = paragraph.add_run(p[2:-2])
            run.font.bold = True
        else:
            run = paragraph.add_run(p)
        set_run_font(run, size=size, color=color)


def add_bookmark(paragraph, name, bmid):
    run = paragraph.add_run()
    bs = OxmlElement('w:bookmarkStart')
    bs.set(qn('w:id'), str(bmid)); bs.set(qn('w:name'), name)
    run._r.append(bs)
    run2 = paragraph.add_run()
    be = OxmlElement('w:bookmarkEnd')
    be.set(qn('w:id'), str(bmid))
    run2._r.append(be)


def add_hyperlink(paragraph, anchor, text, color=MAIN):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle'); rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '%02X%02X%02X' % (color[0], color[1], color[2]))
    rPr.append(color_el)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_h1(doc, text, bmid=None, bookmark_name=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
    if bmid is not None and bookmark_name:
        add_bookmark(p, bookmark_name, bmid)
    run = p.add_run(text)
    set_run_font(run, size=18, bold=True, color=MAIN)
    # 装饰底线（精致风）
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), 'B5838D')
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=15, bold=True, color=SUB)
    return p


def render_normal(doc, text, indent=False, bullet=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    if bullet:
        text = re.sub(r'^-\s+', '• ', text)
    add_inline(p, text)
    return p


def render_highlight(doc, text, color):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    add_inline(p, text, color=BODY)
    shade_para(p, color)
    return p


def render_quote(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.9)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    add_inline(p, text)
    if color:
        shade_para(p, color)
    else:
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '6'); left.set(qn('w:color'), HR)
        pbdr.append(left); pPr.append(pbdr)
    return p


def render_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), HR)
    pbdr.append(bottom); pPr.append(pbdr)


def set_table_borders(table, color="D8C3BC", sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)


def set_col_widths(table, header, doc):
    sec = doc.sections[0]
    usable = sec.page_width - sec.left_margin - sec.right_margin
    ncol = len(header)
    base = 1.3
    weights = [base] * ncol
    wide_terms = ['建议', '风险', '加分', '原因', '分析', '说明', '评价']
    for j, h in enumerate(header):
        for t in wide_terms:
            if t in h:
                weights[j] = 2.4
                break
    total = sum(weights)
    for j, col in enumerate(table.columns):
        w = int(usable * (weights[j] / total))
        col.width = w
        for cell in col.cells:
            cell.width = w


def render_table(doc, rows):
    data = []
    for r in rows:
        cells = [c.strip() for c in r.strip().strip('|').split('|')]
        data.append(cells)
    header = data[0]
    if len(data) > 1 and all(c and set(c) <= set('-: ') for c in data[1]):
        body = data[2:]
    else:
        body = data[1:]
    ncol = len(header)
    table = doc.add_table(rows=1, cols=ncol)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    set_table_borders(table, TBL_BORDER, 4)
    # 表头：实心玫瑰底 + 白字加粗
    hdr = table.rows[0].cells
    for j, c in enumerate(header):
        hdr[j].text = ''
        para = hdr[j].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(para, c, size=10.5, color=TBL_HDR_TXT)
        for run in para.runs:
            run.font.bold = True
        shade_cell(hdr[j], TBL_HDR_FILL)
    # 表体：斑马纹
    for ri, row in enumerate(body):
        cells = table.add_row().cells
        zebra = TBL_ZEBRA if ri % 2 == 0 else TBL_ZEBRA_ALT
        for j in range(ncol):
            val = row[j] if j < len(row) else ''
            cells[j].text = ''
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, val, size=10)
            shade_cell(cells[j], zebra)
    set_col_widths(table, header, doc)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)


def icon_for(text):
    for k, v in SECTION_ICON.items():
        if k in text:
            return v
    return '🌸'


def build_cover(doc, lines):
    company = role = researcher = date = info_date = None
    for ln in lines:
        s = ln.strip()
        m = re.match(r'^\*\*(.+?)\*\*$', s)
        if not m:
            continue
        val = m.group(1)
        if '岗位' in val:
            role = val
        elif '调研人' in val:
            researcher = val
        elif '报告日期' in val:
            date = val
        elif '信息截止' in val:
            info_date = val
        elif company is None:
            company = val
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(company or '比亚迪'), size=22, bold=True, color=MAIN)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run('面试前背调分析报告'), size=15, bold=True, color=SUB)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(role or ''), size=11, color=BODY)
    render_hr(doc)
    for label in [researcher, date, info_date]:
        if label:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(label), size=10.5, color=BODY)
    doc.add_page_break()


def build_toc(doc, toc_items):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run('📑 目 录'), size=15, bold=True, color=MAIN)
    p.paragraph_format.space_after = Pt(10)
    for idx, (text, anchor) in enumerate(toc_items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(6)
        icon = icon_for(text)
        lead = p.add_run(f'{icon}  ')
        set_run_font(lead, size=11, color=SUB)
        add_hyperlink(p, anchor, text, color=MAIN)
        # dotted leader
        dot = p.add_run(' ' + '·' * 6)
        set_run_font(dot, size=10, color=RGBColor.from_string(TOC_LINE))
    doc.add_page_break()


def collect_toc(markdown_lines, body_start):
    items = []
    for ln in markdown_lines[body_start:]:
        s = ln.strip()
        m = re.match(r'^#\s+(.+)$', s)
        if not m:
            continue
        t = m.group(1).strip()
        if t.startswith('【') or t == '信息缺口清单':
            items.append(t)
    return items


def main():
    if len(sys.argv) < 3:
        print("用法: python gen_report_docx.py <输入.md> <输出.docx>")
        sys.exit(1)
    SRC, OUT = sys.argv[1], sys.argv[2]
    with open(SRC, encoding='utf-8') as f:
        lines = f.read().split('\n')

    body_start = 0
    for i, ln in enumerate(lines):
        if ln.strip().startswith('# 【第一部分】'):
            body_start = i
            break
    cover_lines = lines[:body_start]

    # 目录条目（含书签名）
    toc_texts = collect_toc(lines, body_start)
    toc_with_anchor = [(t, 'sec_%d' % i) for i, t in enumerate(toc_texts)]

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(3.0)
    sec.different_first_page_header_footer = True
    hdr = sec.first_page_header.paragraphs[0]
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(hdr.add_run('📋 内部使用・面试调研报告'), size=9, color=GREY)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(10.5)
    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), CN_FONT)
    style.paragraph_format.line_spacing = 1.25

    build_cover(doc, cover_lines)
    build_toc(doc, toc_with_anchor)

    bmid = 0
    i = body_start
    n = len(lines)
    anchor_map = {t: a for t, a in toc_with_anchor}
    while i < n:
        line = lines[i]
        s = line.strip()
        if s.startswith('|'):
            tbl = []
            while i < n and lines[i].strip().startswith('|'):
                tbl.append(lines[i].strip()); i += 1
            render_table(doc, tbl); continue
        if not s:
            i += 1; continue
        if s.startswith('# '):
            t = s[2:].strip()
            anchor = anchor_map.get(t)
            if anchor:
                add_h1(doc, t, bmid=bmid, bookmark_name=anchor); bmid += 1
            else:
                add_h1(doc, t)
            i += 1; continue
        if s.startswith('## '):
            add_h2(doc, s[3:]); i += 1; continue
        if s == '---':
            render_hr(doc); i += 1; continue
        if s.startswith('>'):
            inner = s[1:].lstrip()
            color = None
            for em, col in HL.items():
                if inner.startswith(em):
                    color = col; break
            render_quote(doc, inner, color); i += 1; continue
        matched = False
        for em, col in HL.items():
            if s.startswith(em):
                render_highlight(doc, s, col); matched = True; break
        if matched:
            i += 1; continue
        if re.match(r'^\d+\.\s', s):
            render_normal(doc, s, indent=True); i += 1; continue
        if s.startswith('- '):
            render_normal(doc, s, indent=True, bullet=True); i += 1; continue
        render_normal(doc, s); i += 1

    doc.save(OUT)
    print("SAVED:", OUT, "| 目录条目:", len(toc_with_anchor))


if __name__ == '__main__':
    main()
